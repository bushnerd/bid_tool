import os
import shutil
import logging
from docx import Document
from docx.shared import Inches, Cm
from pdf2image import convert_from_path
from PIL import Image
import tkinter as tk
from tkinter import filedialog

# 设置日志配置
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# 定义路径
root = tk.Tk()
root.withdraw()  # 隐藏主窗口
markdown_file = filedialog.askopenfilename(filetypes=[("Markdown files", "*.md")], title="请选择 Markdown 文件")

# 获取 .md 文件所在目录
md_file_dir = os.path.dirname(markdown_file)

# 获取 .py 文件所在目录
py_file_dir = os.path.dirname(__file__)

source_dir = os.path.join(md_file_dir, 'source_dir')
image_dir = os.path.join(md_file_dir, 'image_dir')
template_path = os.path.join(py_file_dir, '9outline.docx')

# 获取 markdown 文件名（不包括扩展名）
markdown_filename = os.path.splitext(os.path.basename(markdown_file))[0]
output_docx = os.path.join(md_file_dir, f'{markdown_filename}.docx')

# 根据 markdown 结构创建目录
def create_dirs_from_markdown(markdown_file, base_dir):
    """
    Creates directories based on the markdown file structure.

    Args:
        markdown_file (str): The path to the markdown file.
        base_dir (str): The base directory where the directories will be created.

    Returns:
        None

    This function reads the lines of the markdown file and creates directories based on the heading levels specified in the markdown file. The function uses a stack-based approach to store the current paths and ensures that the current paths match the markdown's heading levels. The function creates a new directory based on the current path and the folder name extracted from the markdown line. The function logs the creation of each directory.

    Example:
        create_dirs_from_markdown('path/to/markdown.md', 'path/to/base_dir')
    """
    logging.info(f"Creating directories from markdown file: {markdown_file}")
    logging.info(f"Base directory: {base_dir}")
    with open(markdown_file, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    current_paths = [base_dir]  # 栈结构来存储当前路径
    for line in lines:
        if line.startswith('#'):
            level = line.count('#')
            folder_name = line.strip('#').strip()
            
            # 保证current_paths与markdown的层级一致
            while len(current_paths) > level:
                current_paths.pop()
                logging.info(f"Popped path: {current_paths[-1]}")
            
            # 创建新的路径并存入栈中
            current_path = os.path.join(current_paths[-1], folder_name)
            os.makedirs(current_path, exist_ok=True)
            current_paths.append(current_path)
            logging.info(f"Created directory: {current_path}")

create_dirs_from_markdown(markdown_file, source_dir)

# 等待用户确认
input("请将相应的文件放入创建的目录中，完成后按回车键继续...")

# 创建image目录
if not os.path.exists(image_dir):
    os.makedirs(image_dir)
    logging.info(f'Created image directory: {image_dir}')
else:
    logging.info(f'Image directory already exists: {image_dir}')

# 复制文件夹到image目录
for root, dirs, _ in os.walk(source_dir):
    for dir_name in dirs:
        src_dir = os.path.join(root, dir_name)
        dst_dir = os.path.join(image_dir, os.path.relpath(src_dir, source_dir))
        os.makedirs(dst_dir, exist_ok=True)
        logging.info(f'Created directory: {dst_dir}')

# 导出 PDF 图片
for root, _, files in os.walk(source_dir):
    for file in files:
        file_path = os.path.join(root, file)
        rel_file_path = os.path.relpath(file_path, source_dir)
        file_extension = os.path.splitext(file)[1].lower()

        image_paths = []
        if file_extension == '.pdf':
            logging.info(f'Converting PDF to images: {file_path}')
            images = convert_from_path(file_path, dpi=72, size=(1800, None))
            for i, image in enumerate(images):
                image_path = os.path.join(image_dir, rel_file_path.replace('.pdf', f'_{i}.png'))
                image.save(image_path, 'PNG')
                image_paths.append(image_path)
                logging.info(f'Saved image: {image_path}')
        elif file_extension in ['.png', '.jpg', '.jpeg']:
            image_path = os.path.join(image_dir, rel_file_path)
            shutil.copy(file_path, image_path)
            image_paths = [image_path]
            logging.info(f'Copied image: {image_path}')

# 创建新的docx文档并使用模板样式
doc = Document(template_path)
logging.info(f'Loaded template document: {template_path}')

# 最大宽度设置为14cm
max_width_cm = 14.0
max_width_in = Cm(max_width_cm).inches

# 根据 markdown 结构创建 docx
def create_docx_from_markdown(markdown_file, doc, base_dir, image_dir):
    with open(markdown_file, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    current_paths = [base_dir]  # 栈结构来存储当前路径
    for line in lines:
        if line.startswith('#'):
            level = line.count('#')
            title = line.strip('#').strip()
            while len(current_paths) > level:
                current_paths.pop()
            current_paths.append(title)
            doc.add_heading(title, level=level)
            logging.info(f'Added heading: {title} at level {level}')

            # 获取对应的目录路径
            current_dir = os.path.join(*current_paths)
            image_subdir = os.path.join(image_dir, os.path.relpath(current_dir, base_dir))

            # 检查并插入图片，只插入当前目录的图片
            if os.path.exists(image_subdir):
                for file in os.listdir(image_subdir):
                    image_path = os.path.join(image_subdir, file)
                    if os.path.isfile(image_path):  # 确保只处理文件，不处理子目录
                        with Image.open(image_path) as img:
                            width_in_inches = img.width / img.info.get('dpi', (72, 72))[0]
                            if width_in_inches > max_width_in:
                                width_in_inches = max_width_in
                                logging.info(f'Resizing image: {image_path} to width {width_in_inches} inches')
                            else:
                                logging.info(f'Using original size for image: {image_path}')

                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run()
                        run.add_picture(image_path, width=Inches(width_in_inches))
                        paragraph.style = doc.styles['图片样式']
                        logging.info(f'Inserted image into document with style "图片样式": {image_path}')

                        # 如果是最后一张图片，插入分页符
                        if file == os.listdir(image_subdir)[-1]:
                            doc.add_page_break()
                            logging.info(f'Inserted page break after last image in section: {title}')
# 调用函数
create_docx_from_markdown(markdown_file, doc, source_dir, image_dir)

# 保存 docx 文档
doc.save(output_docx)
logging.info(f'Saved final document: {output_docx}')
